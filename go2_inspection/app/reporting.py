"""整批巡检结果判定、聚合与 Word 报告生成。"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
from typing import Any, Mapping, Sequence

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .errors import ReportNotReadyError
from .storage import CaptureRepository
from .thermal_analysis import (
    THERMAL_ASSOCIATION_WINDOW_SECONDS,
    THERMAL_COOLDOWN_SECONDS,
    THERMAL_THRESHOLD_C,
    ThermalAnalysis,
    ThermalEvent,
    analyze_thermal_samples,
)


STATUS_NORMAL = "normal"
STATUS_ABNORMAL = "abnormal"
STATUS_REVIEW = "review"
STATUS_NOT_APPLICABLE = "not_applicable"

STATUS_LABELS = {
    STATUS_NORMAL: "正常",
    STATUS_ABNORMAL: "异常",
    STATUS_REVIEW: "需复核",
    STATUS_NOT_APPLICABLE: "不适用",
}

STATUS_PRECEDENCE = {
    STATUS_NOT_APPLICABLE: 0,
    STATUS_NORMAL: 1,
    STATUS_REVIEW: 2,
    STATUS_ABNORMAL: 3,
}

ITEM_DEFINITIONS = (
    ("roller_jam", "托辊卡死检测"),
    ("foreign_object", "皮带机异物检测"),
    ("coal_pile", "堆煤检测"),
    ("inspection_marker", "沿线巡检标牌"),
    ("substation_led_meter", "变电硐室 LED 仪表"),
    ("indicator_red", "红色指示灯"),
    ("indicator_green", "绿色指示灯"),
    ("pump_analog_meters", "水泵硐室仪表"),
    ("gas_ch4", "CH4"),
    ("gas_o2", "O2"),
    ("gas_co", "CO"),
    ("gas_h2s", "H2S"),
    ("equipment_temperature", "设备表面温度检测（备选）"),
)

VISUAL_ITEM_IDS = {
    "foreign_object",
    "coal_pile",
    "inspection_marker",
    "substation_led_meter",
    "indicator_red",
    "indicator_green",
    "pump_analog_meters",
}

CAPTURE_ITEM_IDS = {
    "coal_pile",
    "substation_led_meter",
    "pump_analog_meters",
}

GAS_CHANNELS = {
    "ch4": "CH4",
    "o2": "O2",
    "co": "CO",
    "h2s": "H2S",
}

GAS_NORMAL_STATUSES = {"", "normal", "ok"}
GAS_ALARM_STATUSES = {
    "warning",
    "low_alarm",
    "high_alarm",
    "over_range",
    "stel_alarm",
    "twa_alarm",
}

REASON_LABELS = {
    "coal_detector_not_configured": "煤堆检测模型未配置",
    "coal_field_model_not_trained": "堆煤现场模型尚未训练",
    "digital_meter_model_not_trained": "数字表模型尚未训练",
    "station_image_classifier_not_trained": "工位编号模型尚未训练",
    "station_marker_not_detected": "未检测到编号牌",
    "no_confirmed_station_number_readings": "没有可信的工位编号读数",
    "multi_frame_station_numbers_do_not_agree": "多帧工位编号不一致",
    "no_confirmed_frame_readings": "没有可信的数字表读数",
    "multi_frame_readings_do_not_agree": "多帧数字表读数不一致",
    "analog_reference_missing": "缺少指针表正常参考",
    "pointer_not_reliably_detected": "指针未被可靠检测",
    "normal and abnormal reference images are not available": "缺少正常与异常参考图",
}


@dataclass(frozen=True)
class Assessment:
    """报告中的一条检查结论。"""

    item_id: str
    label: str
    status: str
    result: str
    basis: str
    capture_id: str | None = None
    capture_time: str | None = None
    location: str | None = None
    confidence: float | None = None
    evidence_path: str | None = None
    manually_corrected: bool = False


@dataclass(frozen=True)
class GasSummary:
    channel: str
    label: str
    unit: str
    minimum: float | None
    maximum: float | None
    normal_count: int
    abnormal_count: int
    review_count: int
    first_abnormal_time: str | None
    status: str
    basis: str


@dataclass
class BatchReport:
    batch_id: str
    batch_status: str
    generated_at: str
    capture_time_start: str | None
    capture_time_end: str | None
    capture_total: int
    capture_succeeded: int
    capture_failed: int
    gas_row_count: int
    thermal_frame_count: int
    overall_status: str
    overview: list[Assessment]
    details: list[Assessment]
    gas_summaries: list[GasSummary]
    thermal_summary: Assessment
    quality_issues: list[str] = field(default_factory=list)

    @property
    def status_counts(self) -> dict[str, int]:
        counts = {status: 0 for status in STATUS_LABELS}
        for item in self.overview:
            counts[item.status] += 1
        return counts


def build_batch_report(
    repository: CaptureRepository,
    batch_id: str,
) -> BatchReport:
    """从已持久化批次构造可重复生成的报告数据。"""

    batch = repository.get_offline_batch(batch_id)
    if batch["status"] in {"queued", "running"}:
        raise ReportNotReadyError("offline batch is still processing")

    details: list[Assessment] = []
    quality_issues: list[str] = []
    capture_times: list[str] = []
    captures: list[dict[str, Any]] = []

    for item in batch.get("items") or []:
        capture_id = item.get("capture_id")
        if item.get("status") == "succeeded" and capture_id:
            try:
                capture = repository.get_capture(str(capture_id))
            except Exception as exc:  # 报告应尽量产出，并把数据问题留在报告中。
                quality_issues.append(f"抓拍 {capture_id} 无法读取：{exc}")
                continue
            captures.append(capture)
            capture_time = str(capture.get("capture_time") or "")
            if capture_time:
                capture_times.append(capture_time)
            details.extend(_assess_capture(capture))
            if capture.get("error"):
                quality_issues.append(
                    f"抓拍 {capture_id} 处理失败：{capture['error']}"
                )
            result = capture.get("result") or {}
            for warning in result.get("warnings") or []:
                quality_issues.append(f"抓拍 {capture_id}：{warning}")
        elif item.get("status") == "failed":
            quality_issues.append(
                f"批次文件 {item.get('relative_path') or '未知'} 处理失败："
                f"{item.get('error') or '未提供原因'}"
            )

    if not captures:
        details.extend(_empty_capture_assessments())
    details.extend(_pending_required_assessments())

    sensor_samples = repository.sensor_samples_for_batch(batch_id)
    gas_summaries = summarize_gas_samples(sensor_samples)
    details.extend(_gas_assessments(gas_summaries))
    thermal_analysis = analyze_thermal_samples(sensor_samples, captures)
    thermal_summary = _thermal_assessment(batch, sensor_samples, thermal_analysis)
    details.append(thermal_summary)
    details.extend(_thermal_event_assessments(thermal_analysis))
    details.append(_equipment_temperature_assessment(thermal_analysis))

    for diagnostic in batch.get("diagnostics") or []:
        if not isinstance(diagnostic, Mapping):
            continue
        location = diagnostic.get("path") or diagnostic.get("sample_key")
        prefix = f"{location}：" if location else ""
        quality_issues.append(f"{prefix}{diagnostic.get('message') or '数据校验警告'}")
    if batch.get("error"):
        quality_issues.append(f"批次错误：{batch['error']}")

    manually_corrected = sum(
        1 for capture in captures if capture.get("manually_corrected")
    )
    if manually_corrected:
        quality_issues.append(f"本报告采用 {manually_corrected} 条人工修正后的有效结果")
    quality_issues.append(
        "皮带异物、5类巡检标牌、红绿指示灯及3个水泵仪表尚待现场样本与标定"
    )
    if thermal_analysis.unreadable_count:
        quality_issues.append(
            f"有 {thermal_analysis.unreadable_count} 帧热像温度元数据不可读，需人工复核"
        )

    overview = _aggregate_overview(details)
    overall_status = max(
        (item.status for item in overview),
        key=lambda value: STATUS_PRECEDENCE[value],
        default=STATUS_REVIEW,
    )
    ordered_times = sorted(capture_times)
    return BatchReport(
        batch_id=batch_id,
        batch_status=str(batch.get("status") or "unknown"),
        generated_at=datetime.now().astimezone().isoformat(timespec="seconds"),
        capture_time_start=ordered_times[0] if ordered_times else None,
        capture_time_end=ordered_times[-1] if ordered_times else None,
        capture_total=int(batch.get("capture_total") or 0),
        capture_succeeded=int(batch.get("capture_succeeded") or 0),
        capture_failed=int(batch.get("capture_failed") or 0),
        gas_row_count=int(batch.get("gas_row_count") or 0),
        thermal_frame_count=int(batch.get("thermal_frame_count") or 0),
        overall_status=overall_status,
        overview=overview,
        details=details,
        gas_summaries=gas_summaries,
        thermal_summary=thermal_summary,
        quality_issues=_deduplicate(quality_issues),
    )


def summarize_gas_samples(samples: Sequence[Mapping[str, Any]]) -> list[GasSummary]:
    summaries: list[GasSummary] = []
    for channel, label in GAS_CHANNELS.items():
        values: list[float] = []
        units: list[str] = []
        statuses: list[tuple[str, str]] = []
        for sample in samples:
            value = _as_float(sample.get(f"{channel}_value"))
            if value is not None:
                values.append(value)
            unit = str(sample.get(f"{channel}_unit") or "").strip()
            if unit:
                units.append(unit)
            statuses.append(
                (
                    _classify_gas_sample(sample, channel),
                    str(sample.get("captured_at") or ""),
                )
            )
        counts = {
            status: sum(1 for current, _ in statuses if current == status)
            for status in (STATUS_NORMAL, STATUS_ABNORMAL, STATUS_REVIEW)
        }
        first_abnormal = next(
            (time for status, time in statuses if status == STATUS_ABNORMAL),
            None,
        )
        status = (
            STATUS_ABNORMAL
            if counts[STATUS_ABNORMAL]
            else STATUS_REVIEW
            if counts[STATUS_REVIEW] or not statuses
            else STATUS_NORMAL
        )
        basis = (
            "检测到气体报警状态"
            if status == STATUS_ABNORMAL
            else "存在通信故障、缺失或不可读数据"
            if status == STATUS_REVIEW
            else "全部样本状态均为 normal/ok"
        )
        summaries.append(
            GasSummary(
                channel=channel,
                label=label,
                unit=units[0] if units else "-",
                minimum=min(values) if values else None,
                maximum=max(values) if values else None,
                normal_count=counts[STATUS_NORMAL],
                abnormal_count=counts[STATUS_ABNORMAL],
                review_count=counts[STATUS_REVIEW],
                first_abnormal_time=first_abnormal,
                status=status,
                basis=basis,
            )
        )
    return summaries


def render_report_docx(report: BatchReport) -> bytes:
    """按 standard_business_brief 版式生成 Word 文档。"""

    document = Document()
    _configure_document(document, report)
    _add_title_block(document, report)
    _add_overview_section(document, report)
    _add_exception_section(document, report)
    _add_capture_detail_section(document, report)
    _add_sensor_section(document, report)
    _add_quality_section(document, report)
    _add_signoff_section(document)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_prototype_report(evidence_path: str | None = None) -> BatchReport:
    """生成同时覆盖正常、异常、需复核和待训练状态的评审样例。"""

    details = [
        Assessment(
            "roller_jam", "托辊卡死检测", STATUS_ABNORMAL,
            "已识别 3/3 处；最高温 72.30℃",
            "热像最高温严格超过 65.00℃，并关联到 3 个不同编号位置",
            "thermal_demo_001",
            "2026-08-05T09:15:12+08:00", "08号区段；map (1.20, 2.40, 90.0°)",
            evidence_path=evidence_path,
        ),
        Assessment(
            "foreign_object", "皮带机异物检测", STATUS_REVIEW,
            "待样本/不可用", "彩色布条现场正负样本尚未提供",
        ),
        Assessment(
            "coal_pile", "堆煤检测", STATUS_REVIEW,
            "待样本/不可用", "现有样本不足，煤堆检测模型未配置",
        ),
        Assessment(
            "inspection_marker", "沿线巡检标牌", STATUS_REVIEW,
            "待样本/不可用；已识别 0/5 个", "5 类巡检标牌现场样本尚未提供",
        ),
        Assessment(
            "substation_led_meter", "变电硐室 LED 仪表", STATUS_NORMAL,
            "读数 4.38", "LED 数字仪表读数已确认", "demo_capture_001",
            "2026-08-05T09:15:12+08:00", "08号区段", 0.91,
        ),
        Assessment(
            "indicator_red", "红色指示灯", STATUS_REVIEW,
            "待样本/不可用", "红色指示灯现场样本尚未提供",
        ),
        Assessment(
            "indicator_green", "绿色指示灯", STATUS_REVIEW,
            "待样本/不可用", "绿色指示灯现场样本尚未提供",
        ),
        Assessment(
            "pump_analog_meters", "水泵硐室仪表", STATUS_REVIEW,
            "待样本/不可用；已读取 0/3 个", "3 个仪表的现场样本、读数真值和刻度标定尚未提供",
        ),
    ]
    gas_summaries = [
        GasSummary("ch4", "CH4", "%LEL", 0.0, 1.2, 24, 0, 0, None, STATUS_NORMAL, "全部样本状态均为 normal/ok"),
        GasSummary("o2", "O2", "%VOL", 18.9, 20.9, 22, 2, 0, "2026-08-05T09:16:30+08:00", STATUS_ABNORMAL, "检测到 low_alarm"),
        GasSummary("co", "CO", "ppm", 0.0, 3.0, 24, 0, 0, None, STATUS_NORMAL, "全部样本状态均为 normal/ok"),
        GasSummary("h2s", "H2S", "ppm", 0.0, 0.0, 22, 0, 2, None, STATUS_REVIEW, "存在通信超时样本"),
    ]
    details.extend(_gas_assessments(gas_summaries))
    thermal = details[0]
    details.append(
        Assessment(
            "equipment_temperature", "设备表面温度检测（备选）", STATUS_REVIEW,
            "待样本/不可用；热像采集链路已保留", "开关柜测温区域、阈值与现场样本尚未提供",
        )
    )
    overview = _aggregate_overview(details)
    return BatchReport(
        batch_id="prototype-20260805",
        batch_status="completed_with_errors",
        generated_at="2026-08-05T10:00:00+08:00",
        capture_time_start="2026-08-05T09:15:12+08:00",
        capture_time_end="2026-08-05T09:35:45+08:00",
        capture_total=12,
        capture_succeeded=12,
        capture_failed=0,
        gas_row_count=24,
        thermal_frame_count=24,
        overall_status=STATUS_ABNORMAL,
        overview=overview,
        details=details,
        gas_summaries=gas_summaries,
        thermal_summary=thermal,
        quality_issues=[
            "皮带异物、5类巡检标牌、红绿指示灯及3个水泵仪表尚待现场样本与标定",
            "H2S 有 2 条通信超时样本，需人工复核传感器连接",
        ],
    )


def _assess_capture(capture: Mapping[str, Any]) -> list[Assessment]:
    result = capture.get("result") if isinstance(capture.get("result"), Mapping) else {}
    objects = [item for item in result.get("objects") or [] if isinstance(item, Mapping)]
    modules = result.get("modules") if isinstance(result.get("modules"), Mapping) else {}
    capture_id = str(capture.get("capture_id") or "")
    capture_time = str(capture.get("capture_time") or "")
    location = _capture_location(capture, objects, modules)
    evidence = _capture_evidence_path(capture, result)
    manual = bool(capture.get("manually_corrected"))

    capture_assessments: list[Assessment] = []
    foreign_objects = [
        item for item in objects if item.get("type") == "foreign_object"
    ]
    if foreign_objects:
        capture_assessments.append(
            Assessment(
                "foreign_object", "皮带机异物检测", STATUS_ABNORMAL,
                f"检测到异物 {len(foreign_objects)} 处", "皮带机上检测到异物目标",
                capture_id, capture_time, location,
                _max_confidence(foreign_objects), evidence, manual,
            )
        )

    coal_objects = [item for item in objects if item.get("type") == "coal_pile"]
    coal_module = modules.get("coal_presence") if isinstance(modules, Mapping) else None
    if coal_objects or (
        isinstance(coal_module, Mapping)
        and coal_module.get("status") == "confirmed"
        and coal_module.get("present") is True
        and not manual
    ):
        coal = Assessment(
            "coal_pile", "堆煤检测", STATUS_ABNORMAL, "检测到堆煤",
            "检测到煤堆即判为异常", capture_id, capture_time, location,
            _max_confidence(coal_objects), evidence, manual,
        )
    elif manual or (
        isinstance(coal_module, Mapping)
        and coal_module.get("status") == "confirmed"
        and coal_module.get("present") is False
    ):
        coal = Assessment(
            "coal_pile", "堆煤检测", STATUS_NORMAL, "未检测到堆煤",
            "检测模型已运行且 present=false", capture_id, capture_time, location,
            manually_corrected=manual,
        )
    else:
        coal = Assessment(
            "coal_pile", "堆煤检测", STATUS_REVIEW, "待样本/不可用",
            _module_reason(coal_module, "煤堆检测模型未配置或未运行"),
            capture_id, capture_time, location, manually_corrected=manual,
        )

    digital_module = modules.get("digital_meter") if isinstance(modules, Mapping) else None
    digital = _digital_assessment(
        digital_module, objects, capture_id, capture_time, location, evidence, manual
    )
    analog_module = modules.get("analog_meter") if isinstance(modules, Mapping) else None
    analog = _analog_assessment(
        analog_module, objects, capture_id, capture_time, location, evidence, manual
    )
    return [*capture_assessments, coal, digital, analog]


def _digital_assessment(
    module: object,
    objects: Sequence[Mapping[str, Any]],
    capture_id: str,
    capture_time: str,
    location: str,
    evidence: str | None,
    manual: bool,
) -> Assessment:
    value: object | None = None
    confidence: float | None = None
    if isinstance(module, Mapping) and module.get("status") == "confirmed":
        value = module.get("raw_text")
        if value is None:
            value = module.get("value")
        confidence = _as_float(module.get("confidence"))
    if value is None:
        meters = [item for item in objects if item.get("type") == "digital_meter"]
        for meter in meters:
            value = meter.get("raw_text")
            if value is None:
                value = meter.get("value")
            if value is not None:
                confidence = _as_float(
                    meter.get("reading_confidence") or meter.get("confidence")
                )
                break
    result = f"读数 {value}" if value is not None else "未获得可信数字表读数"
    basis = (
        "LED 数字仪表读数已确认"
        if value is not None
        else _module_reason(module, "数字表不可读、未检测到或模型未运行")
    )
    return Assessment(
        "substation_led_meter", "变电硐室 LED 仪表",
        STATUS_NORMAL if value is not None else STATUS_REVIEW, result, basis,
        capture_id, capture_time, location, confidence,
        evidence if value is not None else None, manual,
    )


def _analog_assessment(
    module: object,
    objects: Sequence[Mapping[str, Any]],
    capture_id: str,
    capture_time: str,
    location: str,
    evidence: str | None,
    manual: bool,
) -> Assessment:
    meters = [item for item in objects if item.get("type") == "analog_meter"]
    if not meters and isinstance(module, Mapping):
        meters = [
            item for item in module.get("meters") or [] if isinstance(item, Mapping)
        ]
    readings = [
        value
        for item in meters
        for value in (_meter_display_value(item),)
        if value is not None
    ]
    complete = len(readings) == 3
    status = STATUS_NORMAL if complete else STATUS_REVIEW
    result = (
        f"已读取 3/3 个：{'、'.join(readings)}"
        if complete
        else f"待样本/不可用；已读取 {len(readings)}/3 个"
    )
    basis = (
        "3 个水泵硐室仪表读数均已确认"
        if complete
        else "3 个仪表的现场样本、读数真值和刻度标定尚未完整提供"
    )
    return Assessment(
        "pump_analog_meters", "水泵硐室仪表", status, result, basis,
        capture_id, capture_time, location, _max_meter_confidence(meters),
        evidence if meters else None, manual,
    )


def _empty_capture_assessments() -> list[Assessment]:
    return [
        Assessment(item_id, label, STATUS_REVIEW, "无可用抓拍结果", "批次没有成功处理的可见光抓拍")
        for item_id, label in ITEM_DEFINITIONS
        if item_id in CAPTURE_ITEM_IDS
    ]


def _pending_required_assessments() -> list[Assessment]:
    return [
        Assessment(
            "foreign_object", "皮带机异物检测", STATUS_REVIEW,
            "待样本/不可用", "彩色布条现场正负样本尚未提供",
        ),
        Assessment(
            "inspection_marker", "沿线巡检标牌", STATUS_REVIEW,
            "待样本/不可用；已识别 0/5 个", "5 类巡检标牌现场样本尚未提供；编号牌不能替代巡检标牌",
        ),
        Assessment(
            "indicator_red", "红色指示灯", STATUS_REVIEW,
            "待样本/不可用", "红色指示灯现场样本尚未提供",
        ),
        Assessment(
            "indicator_green", "绿色指示灯", STATUS_REVIEW,
            "待样本/不可用", "绿色指示灯现场样本尚未提供",
        ),
    ]


def _gas_assessments(summaries: Sequence[GasSummary]) -> list[Assessment]:
    result: list[Assessment] = []
    for item in summaries:
        value_range = (
            f"{_format_number(item.minimum)}–{_format_number(item.maximum)} {item.unit}"
            if item.minimum is not None and item.maximum is not None
            else "无有效读数"
        )
        counts = (
            f"正常 {item.normal_count}，异常 {item.abnormal_count}，"
            f"需复核 {item.review_count}"
        )
        result.append(
            Assessment(
                f"gas_{item.channel}", item.label, item.status,
                f"范围 {value_range}；{counts}", item.basis,
                capture_time=item.first_abnormal_time,
            )
        )
    return result


def _thermal_assessment(
    batch: Mapping[str, Any],
    samples: Sequence[Mapping[str, Any]],
    analysis: ThermalAnalysis,
) -> Assessment:
    recorded = int(batch.get("thermal_frame_count") or 0)
    stored = [
        Path(str(sample["thermal_stored_path"]))
        for sample in samples
        if sample.get("thermal_stored_path")
    ]
    valid = sum(1 for path in stored if path.is_file())
    missing = max(0, recorded - valid)
    position_events = _hottest_event_per_position(analysis)
    unknown_events = [
        event for event in analysis.events if event.station_number is None
    ]
    result = (
        f"已登记 {recorded} 帧；可访问 {valid} 帧；"
        f"温度可读 {analysis.readable_count} 帧"
    )
    if analysis.maximum_c is not None:
        result += f"；批次最高温 {analysis.maximum_c:.2f}℃"
    result += (
        f"；超温候选 {len(analysis.candidates)} 条；"
        f"10秒抑制 {analysis.suppressed_count} 条；"
        f"已识别 {len(position_events)}/3 处"
    )
    if unknown_events:
        result += f"；位置未知候选 {len(unknown_events)} 条"
    if missing:
        result += f"；缺失 {missing} 帧"
    if position_events:
        status = STATUS_ABNORMAL
        basis = (
            f"检测到最高温严格超过 {THERMAL_THRESHOLD_C:.2f}℃ 的热像；"
            f"同编号异常在 {THERMAL_COOLDOWN_SECONDS:.0f} 秒内抑制重复，"
            "报告按编号位置保留最高温证据"
        )
    elif unknown_events:
        status = STATUS_REVIEW
        basis = (
            f"检测到最高温严格超过 {THERMAL_THRESHOLD_C:.2f}℃ 的热像，"
            "但未关联到可信编号位置"
        )
    elif (
        not recorded
        or not analysis.readable_count
        or analysis.unreadable_count
        or missing
    ):
        status = STATUS_REVIEW
        basis = (
            "未发现可用于分析的红外热像"
            if not recorded
            else "存在温度元数据不可读或归档文件缺失"
        )
    else:
        status = STATUS_NORMAL
        basis = f"全部可读热像最高温均不超过 {THERMAL_THRESHOLD_C:.2f}℃"
    return Assessment("roller_jam", "托辊卡死检测", status, result, basis)


def _thermal_event_assessments(
    analysis: ThermalAnalysis,
) -> list[Assessment]:
    details: list[Assessment] = []
    events = [
        *_hottest_event_per_position(analysis),
        *(event for event in analysis.events if event.station_number is None),
    ]
    for event in events:
        if event.station_number is None:
            station_text = "编号未知"
            location = "编号未知（需复核）"
            status = STATUS_REVIEW
            association = (
                f"前后 {THERMAL_ASSOCIATION_WINDOW_SECONDS:.0f} 秒内"
                "没有可信编号识别结果"
            )
        else:
            station_text = f"关联 {event.station_number} 号编号牌"
            location = f"{event.station_number} 号牌位置"
            status = STATUS_ABNORMAL
            association = (
                f"最近编号帧 {event.station_capture_id}，"
                f"时间差 {event.association_delta_seconds or 0.0:.3f} 秒"
            )
        image_path = event.image_path
        capture_id = Path(image_path).stem if image_path else event.sample_key
        details.append(
            Assessment(
                "roller_jam",
                "疑似托辊卡死",
                status,
                f"最高温 {event.maximum_c:.2f}℃；{station_text}",
                (
                    f"最高温严格超过 {THERMAL_THRESHOLD_C:.2f}℃；"
                    f"{association}"
                ),
                capture_id=capture_id,
                capture_time=event.captured_at,
                location=location,
                evidence_path=image_path,
            )
        )
    return details


def _hottest_event_per_position(
    analysis: ThermalAnalysis,
) -> list[ThermalEvent]:
    by_position: dict[int, ThermalEvent] = {}
    for event in analysis.candidates:
        if event.station_number is None:
            continue
        current = by_position.get(event.station_number)
        if current is None or (event.maximum_c, event.captured_at) > (
            current.maximum_c,
            current.captured_at,
        ):
            by_position[event.station_number] = event
    return [by_position[number] for number in sorted(by_position)]


def _equipment_temperature_assessment(
    analysis: ThermalAnalysis,
) -> Assessment:
    if analysis.readable_count:
        result = f"热像采集链路可用；已读取 {analysis.readable_count} 帧"
        if analysis.maximum_c is not None:
            result += f"；整帧最高温 {analysis.maximum_c:.2f}℃"
    else:
        result = "没有可用于设备表面测温的有效热像"
    return Assessment(
        "equipment_temperature",
        "设备表面温度检测（备选）",
        STATUS_REVIEW,
        f"待样本/不可用；{result}",
        "开关柜测温区域、异常阈值与现场样本尚未提供，当前不得判定正常或异常",
    )


def _aggregate_overview(details: Sequence[Assessment]) -> list[Assessment]:
    overview: list[Assessment] = []
    for item_id, label in ITEM_DEFINITIONS:
        candidates = [item for item in details if item.item_id == item_id]
        if not candidates:
            overview.append(
                Assessment(item_id, label, STATUS_REVIEW, "无可用结果", "该检查项没有可用数据")
            )
            continue
        status = max(
            (item.status for item in candidates),
            key=lambda value: STATUS_PRECEDENCE[value],
        )
        representative = next(
            (item for item in candidates if item.status == status), candidates[0]
        )
        counts = {
            current: sum(1 for item in candidates if item.status == current)
            for current in STATUS_LABELS
        }
        count_text = "，".join(
            f"{STATUS_LABELS[current]} {counts[current]}"
            for current in (STATUS_NORMAL, STATUS_ABNORMAL, STATUS_REVIEW, STATUS_NOT_APPLICABLE)
            if counts[current]
        )
        overview.append(
            Assessment(
                item_id, label, status,
                f"{count_text}；{representative.result}" if count_text else representative.result,
                representative.basis,
            )
        )
    return overview


def _classify_gas_sample(sample: Mapping[str, Any], channel: str) -> str:
    value = _as_float(sample.get(f"{channel}_value"))
    raw_status = str(sample.get(f"{channel}_status") or "").strip().lower()
    error_status = _gas_status_from_error(str(sample.get("gas_error") or ""), channel)
    status = error_status or raw_status
    if status in GAS_ALARM_STATUSES:
        return STATUS_ABNORMAL
    if status in GAS_NORMAL_STATUSES:
        if value is not None:
            return STATUS_NORMAL
        return STATUS_REVIEW
    if status:
        return STATUS_REVIEW
    return STATUS_NORMAL if value is not None else STATUS_REVIEW


def _gas_status_from_error(error: str, channel: str) -> str | None:
    if not error:
        return None
    label = GAS_CHANNELS[channel]
    match = re.search(
        rf"(?i)(?:^|[；;])\s*{re.escape(label)}状态[：:]\s*([a-z_]+)",
        error,
    )
    if match:
        return match.group(1).lower()
    upper = error.upper()
    if label in upper and any(
        marker in error for marker in ("缺少", "读取失败", "超时", "故障")
    ):
        return "read_error"
    return None


def _capture_location(
    capture: Mapping[str, Any],
    objects: Sequence[Mapping[str, Any]],
    modules: Mapping[str, Any],
) -> str:
    station_id = str(capture.get("station_id") or "").strip()
    station_module = modules.get("station_number")
    recognized_number = (
        station_module.get("number")
        if isinstance(station_module, Mapping)
        and station_module.get("status") == "confirmed"
        else None
    )
    location = next(
        (str(item.get("location_text")) for item in objects if item.get("location_text")),
        (
            f"{recognized_number} 号牌位置"
            if recognized_number is not None
            else f"{station_id}号区段" if station_id else "未提供区段"
        ),
    )
    pose = capture.get("robot_pose") if isinstance(capture.get("robot_pose"), Mapping) else {}
    coordinates = [pose.get("x_m"), pose.get("y_m"), pose.get("yaw_deg")]
    if any(value is not None for value in coordinates):
        location += (
            f"；{pose.get('frame') or 'map'} "
            f"({_display_optional(coordinates[0])}, {_display_optional(coordinates[1])}, "
            f"{_display_optional(coordinates[2])}°)"
        )
    return location


def _capture_evidence_path(
    capture: Mapping[str, Any], result: Mapping[str, Any]
) -> str | None:
    annotated = result.get("annotated_image")
    if annotated and Path(str(annotated)).is_file():
        return str(annotated)
    for image in capture.get("images") or []:
        if not isinstance(image, Mapping):
            continue
        stored = image.get("stored_path")
        if stored and Path(str(stored)).is_file():
            return str(stored)
    return None


def _module_reason(module: object, fallback: str) -> str:
    if not isinstance(module, Mapping):
        return fallback
    reason = str(module.get("reason") or "").strip()
    return REASON_LABELS.get(reason, reason or fallback)


def _object_label(item: Mapping[str, Any]) -> str:
    return str(item.get("class_cn") or item.get("class") or item.get("type") or "未知目标")


def _max_confidence(items: Sequence[Mapping[str, Any]]) -> float | None:
    values = [_as_float(item.get("confidence")) for item in items]
    finite = [value for value in values if value is not None]
    return max(finite) if finite else None


def _max_meter_confidence(items: Sequence[Mapping[str, Any]]) -> float | None:
    values = [
        _as_float(item.get("meter_confidence") or item.get("confidence"))
        for item in items
    ]
    finite = [value for value in values if value is not None]
    return max(finite) if finite else None


def _meter_display_value(item: Mapping[str, Any]) -> str | None:
    for field in ("raw_text", "value", "reading"):
        value = item.get(field)
        if value is not None and str(value).strip():
            return str(value)
    return None


def _as_float(value: object) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _display_optional(value: object) -> str:
    number = _as_float(value)
    return "-" if number is None else f"{number:.2f}"


def _format_number(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value:.3f}".rstrip("0").rstrip(".")


def _deduplicate(values: Sequence[str]) -> list[str]:
    return list(dict.fromkeys(value for value in values if value))


# --- Word rendering helpers -------------------------------------------------

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM_DXA = 80
CELL_MARGIN_LEFT_RIGHT_DXA = 120

COLOR_HEADING = "2E74B5"
COLOR_HEADING_DARK = "1F4D78"
COLOR_MUTED = "5B6573"
COLOR_TABLE_HEADER = "F2F4F7"
STATUS_FILLS = {
    STATUS_NORMAL: "E8F3EC",
    STATUS_ABNORMAL: "FCE8E6",
    STATUS_REVIEW: "FFF4D6",
    STATUS_NOT_APPLICABLE: "ECEFF2",
}
STATUS_TEXT = {
    STATUS_NORMAL: "246B3B",
    STATUS_ABNORMAL: "9B1C1C",
    STATUS_REVIEW: "7A5A00",
    STATUS_NOT_APPLICABLE: "5B6573",
}


def _configure_document(document: Document, report: BatchReport) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, COLOR_HEADING, 16, 8),
        ("Heading 2", 13, COLOR_HEADING, 12, 6),
        ("Heading 3", 12, COLOR_HEADING_DARK, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"GO2 智能巡检 | 批次 {report.batch_id}")
    _set_run_font(run, 9, COLOR_MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("第 ")
    _set_run_font(run, 9, COLOR_MUTED)
    _add_field(footer, "PAGE")
    run = footer.add_run(" 页 / 共 ")
    _set_run_font(run, 9, COLOR_MUTED)
    _add_field(footer, "NUMPAGES")
    run = footer.add_run(" 页")
    _set_run_font(run, 9, COLOR_MUTED)

    properties = document.core_properties
    properties.title = "煤矿实验室 GO2 智能巡检报告"
    properties.subject = f"批次 {report.batch_id} 巡检结果"
    properties.author = "GO2 Inspection"


def _add_title_block(document: Document, report: BatchReport) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("煤矿实验室 GO2 智能巡检报告")
    _set_run_font(run, 24, "0B2545", bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run("整批离线数据识别、异常汇总与复核记录")
    _set_run_font(run, 12, COLOR_MUTED)

    metadata = (
        ("批次编号", report.batch_id),
        ("巡检时间", _time_range(report.capture_time_start, report.capture_time_end)),
        ("报告生成", _display_time(report.generated_at)),
        ("数据规模", f"抓拍 {report.capture_total} 组；气体 {report.gas_row_count} 条；热像 {report.thermal_frame_count} 帧"),
        ("批次状态", report.batch_status),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        _set_run_font(label_run, 10.5, "202A35", bold=True)
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, 10.5, "202A35")

    callout = document.add_table(rows=1, cols=1)
    _set_table_geometry(callout, [PAGE_WIDTH_DXA])
    cell = callout.cell(0, 0)
    _set_cell_fill(cell, STATUS_FILLS[report.overall_status])
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"整体结论：{STATUS_LABELS[report.overall_status]}")
    _set_run_font(run, 15, STATUS_TEXT[report.overall_status], bold=True)

    counts = report.status_counts
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(6)
    summary.paragraph_format.space_after = Pt(6)
    run = summary.add_run(
        f"正常 {counts[STATUS_NORMAL]} 项   |   异常 {counts[STATUS_ABNORMAL]} 项   |   "
        f"需复核 {counts[STATUS_REVIEW]} 项   |   不适用 {counts[STATUS_NOT_APPLICABLE]} 项"
    )
    _set_run_font(run, 10, COLOR_MUTED, bold=True)


def _add_overview_section(document: Document, report: BatchReport) -> None:
    document.add_heading("1. 项目总览", level=1)
    rows = []
    for item in report.overview:
        rows.append(
            ([item.label, STATUS_LABELS[item.status], item.result, item.basis], item.status)
        )
    _add_table(
        document,
        ["检查项目", "结论", "识别结果", "判定依据与备注"],
        rows,
        [1500, 1000, 2800, 4060],
        status_column=1,
    )


def _add_exception_section(document: Document, report: BatchReport) -> None:
    document.add_heading("2. 异常与待复核明细", level=1)
    exceptions = [
        item for item in report.details
        if item.status in {STATUS_ABNORMAL, STATUS_REVIEW}
    ]
    if not exceptions:
        paragraph = document.add_paragraph("本批次未检测到明确异常，也没有待复核项目。")
        paragraph.paragraph_format.space_after = Pt(8)
        return
    rows = []
    for item in exceptions:
        time_location = "\n".join(
            value for value in (_display_time(item.capture_time), item.location) if value
        ) or "批次汇总"
        result = item.result
        if item.confidence is not None:
            result += f"（置信度 {item.confidence:.1%}）"
        rows.append(
            ([time_location, item.label, result, STATUS_LABELS[item.status], item.basis], item.status)
        )
    _add_table(
        document,
        ["时间 / 位置", "项目", "检测结果", "结论", "原因 / 处置提示"],
        rows,
        [1700, 1200, 2200, 1000, 3260],
        status_column=3,
    )

    evidence_seen: set[str] = set()
    for item in exceptions:
        if item.status != STATUS_ABNORMAL or not item.capture_id:
            continue
        if item.capture_id in evidence_seen:
            continue
        evidence_seen.add(item.capture_id)
        document.add_heading(f"证据图：{item.capture_id}", level=3)
        paragraph = document.add_paragraph(
            f"{item.label} - {item.result}；{item.basis}。"
        )
        paragraph.paragraph_format.keep_with_next = True
        if not item.evidence_path:
            document.add_paragraph("未找到可嵌入的标注图或原图，请在上位机中复核原始记录。")
            continue
        try:
            stream = _compressed_image(item.evidence_path)
            picture = document.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.paragraph_format.space_after = Pt(8)
            picture.add_run().add_picture(stream, width=Inches(5.8))
        except Exception as exc:
            document.add_paragraph(f"证据图无法嵌入：{exc}")


def _add_capture_detail_section(document: Document, report: BatchReport) -> None:
    document.add_heading("3. 逐抓拍视觉识别明细", level=1)
    visual = [item for item in report.details if item.item_id in VISUAL_ITEM_IDS]
    rows = []
    for item in visual:
        time_location = "\n".join(
            value for value in (_display_time(item.capture_time), item.location) if value
        ) or "无可用抓拍"
        result = item.result
        if item.manually_corrected:
            result += "（人工修正后结果）"
        rows.append(
            ([time_location, item.label, result, STATUS_LABELS[item.status], item.basis], item.status)
        )
    _add_table(
        document,
        ["时间 / 位置", "识别项目", "识别结果", "结论", "说明"],
        rows,
        [1700, 1300, 2300, 1000, 3060],
        status_column=3,
    )


def _add_sensor_section(document: Document, report: BatchReport) -> None:
    document.add_heading("4. 传感器汇总", level=1)
    document.add_heading("4.1 气体监测", level=2)
    rows = []
    for item in report.gas_summaries:
        rows.append(
            (
                [
                    item.label,
                    item.unit,
                    _format_number(item.minimum),
                    _format_number(item.maximum),
                    f"{item.normal_count} / {item.abnormal_count} / {item.review_count}",
                    STATUS_LABELS[item.status],
                    _display_time(item.first_abnormal_time) or item.basis,
                ],
                item.status,
            )
        )
    _add_table(
        document,
        ["气体", "单位", "最小值", "最大值", "正常/异常/复核", "结论", "首次异常 / 备注"],
        rows,
        [850, 800, 850, 850, 1400, 900, 3710],
        status_column=5,
    )

    document.add_heading("4.2 红外热像", level=2)
    thermal = report.thermal_summary
    equipment_temperature = next(
        (
            item for item in report.details
            if item.item_id == "equipment_temperature"
        ),
        Assessment(
            "equipment_temperature", "设备表面温度检测（备选）",
            STATUS_REVIEW, "无可用结果", "开关柜测温尚未配置",
        ),
    )
    _add_table(
        document,
        ["检查项目", "数据状态", "结论", "说明"],
        [
            ([thermal.label, thermal.result, STATUS_LABELS[thermal.status], thermal.basis], thermal.status),
            (
                [
                    equipment_temperature.label,
                    equipment_temperature.result,
                    STATUS_LABELS[equipment_temperature.status],
                    equipment_temperature.basis,
                ],
                equipment_temperature.status,
            ),
        ],
        [1500, 2600, 1100, 4160],
        status_column=2,
    )


def _add_quality_section(document: Document, report: BatchReport) -> None:
    document.add_heading("5. 数据质量与系统状态", level=1)
    issues = report.quality_issues or ["未记录额外的数据质量问题"]
    rows = [([str(index), issue], None) for index, issue in enumerate(issues, 1)]
    _add_table(document, ["序号", "问题 / 说明"], rows, [700, 8660])


def _add_signoff_section(document: Document) -> None:
    document.add_heading("6. 复核与签字", level=1)
    paragraph = document.add_paragraph(
        "复核人应重点确认所有“异常”和“需复核”项目；本报告用于实验室巡检记录，"
        "不能替代煤矿现场法定安全检测或联锁装置。"
    )
    paragraph.paragraph_format.space_after = Pt(10)
    rows = [
        (["复核人", "____________________", "复核日期", "____年__月__日"], None),
        (["处理意见", "", "签字", "____________________"], None),
    ]
    _add_table(document, [], rows, [1200, 3200, 1200, 3760])


def _add_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[tuple[Sequence[str], str | None]],
    widths: Sequence[int],
    *,
    status_column: int | None = None,
) -> None:
    table = document.add_table(rows=1 if headers else 0, cols=len(widths))
    table.style = "Table Grid"
    if headers:
        header = table.rows[0]
        _set_repeat_table_header(header)
        for index, value in enumerate(headers):
            cell = header.cells[index]
            _set_cell_fill(cell, COLOR_TABLE_HEADER)
            _set_cell_text(cell, value, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for values, status in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            align = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index == status_column or len(str(value)) <= 12
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_cell_text(row.cells[index], str(value), size=9, align=align)
        if status is not None and status_column is not None:
            cell = row.cells[status_column]
            _set_cell_fill(cell, STATUS_FILLS[status])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(STATUS_TEXT[status])
                    run.bold = True
    _set_table_geometry(table, widths)
    trailing = document.add_paragraph()
    trailing.paragraph_format.space_before = Pt(0)
    trailing.paragraph_format.space_after = Pt(2)


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError("table column widths must total 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.insert(0, width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(value))
        grid.append(grid_column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[index]))
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (
                ("top", CELL_MARGIN_TOP_BOTTOM_DXA),
                ("bottom", CELL_MARGIN_TOP_BOTTOM_DXA),
                ("start", CELL_MARGIN_LEFT_RIGHT_DXA),
                ("end", CELL_MARGIN_LEFT_RIGHT_DXA),
            ):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    size: float = 9,
    align: Any = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    _set_run_font(run, size, "202A35", bold=bold)


def _set_cell_fill(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_run_font(
    run: Any,
    size: float,
    color: str,
    *,
    bold: bool = False,
) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    _set_run_font(run, 9, COLOR_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, display, end])


def _compressed_image(path: str) -> BytesIO:
    stream = BytesIO()
    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image)
        image = image.convert("RGB")
        image.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
        image.save(stream, format="JPEG", quality=82, optimize=True)
    stream.seek(0)
    return stream


def _display_time(value: str | None) -> str:
    if not value:
        return ""
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
        return parsed.isoformat(sep=" ", timespec="seconds")
    except ValueError:
        return value


def _time_range(start: str | None, end: str | None) -> str:
    if not start and not end:
        return "无可用抓拍时间"
    if start == end or not end:
        return _display_time(start)
    return f"{_display_time(start)} 至 {_display_time(end)}"
