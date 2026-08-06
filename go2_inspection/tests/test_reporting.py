from io import BytesIO
import json
import os
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient

from app.api import create_app
from app.domain import CaptureMetadata, InspectionResult, RobotPose
from app.reporting import (
    STATUS_ABNORMAL,
    STATUS_NORMAL,
    STATUS_REVIEW,
    build_batch_report,
    build_prototype_report,
    render_report_docx,
    summarize_gas_samples,
    _hottest_event_per_position,
    _thermal_assessment,
    _thermal_event_assessments,
)
from app.storage import CaptureRepository
from app.thermal_analysis import analyze_thermal_samples
from tests.test_offline_batches import write_batch
from tests.test_service import MINIMAL_PNG


class ReportJudgementTests(unittest.TestCase):
    def test_thermal_report_deduplicates_positions_and_reviews_unknowns(self) -> None:
        def thermal(timestamp: str, maximum: float, suffix: str) -> dict:
            return {
                "sample_key": f"thermal_{suffix}",
                "captured_at": timestamp,
                "thermal_stored_path": f"missing_{suffix}.png",
                "thermal_maximum_c": maximum,
                "thermal_metadata_status": "valid",
            }

        def numbered(timestamp: str, number: int, suffix: str) -> dict:
            return {
                "capture_id": f"capture_{suffix}",
                "capture_time": timestamp,
                "result": {
                    "modules": {
                        "station_number": {
                            "status": "confirmed",
                            "number": number,
                            "confidence": 0.9,
                        }
                    }
                },
            }

        samples = [
            thermal("2026-08-05T09:00:00+08:00", 70.0, "a"),
            thermal("2026-08-05T09:00:01+08:00", 80.0, "b"),
            thermal("2026-08-05T09:00:12+08:00", 75.0, "c"),
            thermal("2026-08-05T09:00:20+08:00", 76.0, "d"),
            thermal("2026-08-05T09:00:30+08:00", 77.0, "e"),
        ]
        captures = [
            numbered("2026-08-05T09:00:00+08:00", 1, "one"),
            numbered("2026-08-05T09:00:12+08:00", 1, "one_later"),
            numbered("2026-08-05T09:00:20+08:00", 2, "two"),
            numbered("2026-08-05T09:00:30+08:00", 3, "three"),
        ]

        analysis = analyze_thermal_samples(samples, captures)
        hottest = _hottest_event_per_position(analysis)
        summary = _thermal_assessment(
            {"thermal_frame_count": len(samples)}, samples, analysis
        )

        self.assertEqual([item.station_number for item in hottest], [1, 2, 3])
        self.assertEqual(hottest[0].maximum_c, 80.0)
        self.assertIn("已识别 3/3 处", summary.result)
        self.assertEqual(summary.status, STATUS_ABNORMAL)

        unknown_samples = [thermal("2026-08-05T10:00:00+08:00", 70.0, "unknown")]
        unknown = analyze_thermal_samples(unknown_samples, [])
        unknown_summary = _thermal_assessment(
            {"thermal_frame_count": 1}, unknown_samples, unknown
        )
        unknown_details = _thermal_event_assessments(unknown)
        self.assertEqual(unknown_summary.status, STATUS_REVIEW)
        self.assertIn("已识别 0/3 处", unknown_summary.result)
        self.assertEqual(unknown_details[0].status, STATUS_REVIEW)

    def test_gas_summary_distinguishes_alarm_fault_and_normal(self) -> None:
        summaries = summarize_gas_samples(
            [
                {
                    "captured_at": "2026-08-05T09:00:00+08:00",
                    "ch4_value": 1.2,
                    "ch4_unit": "%LEL",
                    "ch4_status": "normal",
                    "o2_value": 19.5,
                    "o2_unit": "%VOL",
                    "o2_status": None,
                    "co_value": 2,
                    "co_unit": "ppm",
                    "co_status": "normal",
                    "h2s_value": None,
                    "h2s_unit": "ppm",
                    "h2s_status": "timeout",
                    "gas_error": "O2状态：low_alarm",
                },
                {
                    "captured_at": "2026-08-05T09:00:05+08:00",
                    "ch4_value": 0.8,
                    "ch4_unit": "%LEL",
                    "ch4_status": "ok",
                    "o2_value": 20.9,
                    "o2_unit": "%VOL",
                    "o2_status": "normal",
                    "co_value": 1,
                    "co_unit": "ppm",
                    "co_status": "ok",
                    "h2s_value": 0,
                    "h2s_unit": "ppm",
                    "h2s_status": "normal",
                    "gas_error": None,
                },
            ]
        )
        by_channel = {item.channel: item for item in summaries}
        self.assertEqual(by_channel["ch4"].status, STATUS_NORMAL)
        self.assertEqual(by_channel["o2"].status, STATUS_ABNORMAL)
        self.assertEqual(by_channel["o2"].first_abnormal_time, "2026-08-05T09:00:00+08:00")
        self.assertEqual(by_channel["h2s"].status, STATUS_REVIEW)

    def test_repository_report_uses_fixed_safety_rules(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source"
            source.mkdir()
            repository = CaptureRepository(root / "inspection.db", root / "runtime")
            capture_id = "report_capture_001"
            repository.create_offline_batch(
                "report-batch",
                source,
                [{"relative_path": "visible/report_capture_001", "capture_id": capture_id}],
                gas_row_count=1,
                thermal_frame_count=1,
            )
            metadata = CaptureMetadata(
                capture_id=capture_id,
                capture_time="2026-08-05T09:15:12+08:00",
                station_id="08",
                robot_pose=RobotPose(frame="map", x_m=1.2, y_m=2.4, yaw_deg=90),
                camera_id="test",
                image_names=("frame.png",),
                batch_id="report-batch",
            )
            image_path = repository.save_capture(
                metadata, [("frame.png", MINIMAL_PNG)], 1024 * 1024
            )[0]
            result = InspectionResult(
                capture_id=capture_id,
                station_id="08",
                capture_pose=metadata.robot_pose,
                objects=[
                    {
                        "type": "analog_meter",
                        "class": "analog_meter",
                        "class_cn": "水泵仪表",
                        "raw_text": "0.42 MPa",
                        "meter_confidence": 0.90,
                        "confidence": 0.95,
                        "bbox_xyxy": [1, 1, 10, 10],
                    },
                    {
                        "type": "analog_meter",
                        "class": "analog_meter",
                        "class_cn": "水泵仪表",
                        "raw_text": "0.46 MPa",
                        "meter_confidence": 0.88,
                        "confidence": 0.93,
                        "bbox_xyxy": [11, 1, 20, 10],
                    },
                    {
                        "type": "analog_meter",
                        "class": "analog_meter",
                        "class_cn": "水泵仪表",
                        "raw_text": "0.44 MPa",
                        "meter_confidence": 0.91,
                        "confidence": 0.94,
                        "bbox_xyxy": [1, 1, 10, 10],
                    },
                ],
                modules={
                    "coal_presence": {"enabled": True, "status": "confirmed", "present": False},
                    "station_number": {"enabled": True, "status": "confirmed", "number": 9, "confidence": 0.93},
                    "digital_meter": {"enabled": True, "status": "confirmed", "raw_text": "067.8", "confidence": 0.91},
                    "analog_meter": {"enabled": True, "status": "confirmed"},
                },
                annotated_image=str(image_path),
                processing_parameters={"detector": {"mode": "gpu"}},
            )
            repository.save_result(result)
            repository.update_offline_item(
                "report-batch",
                "visible/report_capture_001",
                status="succeeded",
                capture_id=capture_id,
            )
            thermal = root / "thermal.png"
            thermal.write_bytes(MINIMAL_PNG)
            repository.replace_sensor_samples(
                "report-batch",
                [
                    {
                        "sample_key": "20260805_091512_000001",
                        "captured_at": "2026-08-05T09:15:12+08:00",
                        "sample_id": "000001",
                        "ch4_value": 0,
                        "ch4_unit": "%LEL",
                        "ch4_status": "normal",
                        "o2_value": 19.5,
                        "o2_unit": "%VOL",
                        "o2_status": "low_alarm",
                        "co_value": 2,
                        "co_unit": "ppm",
                        "co_status": "normal",
                        "h2s_value": 0,
                        "h2s_unit": "ppm",
                        "h2s_status": "normal",
                        "thermal_stored_path": str(thermal),
                        "thermal_minimum_c": 24.0,
                        "thermal_maximum_c": 70.0,
                        "thermal_average_c": 30.0,
                        "thermal_metadata_status": "valid",
                    }
                ],
                gas_row_count=1,
                thermal_frame_count=1,
                diagnostics=[],
            )
            repository.finish_offline_batch("report-batch")

            report = build_batch_report(repository, "report-batch")
            overview = {item.item_id: item for item in report.overview}
            self.assertEqual(len(report.overview), 13)
            self.assertEqual(report.overall_status, STATUS_ABNORMAL)
            self.assertEqual(overview["foreign_object"].status, STATUS_REVIEW)
            self.assertEqual(overview["coal_pile"].status, STATUS_NORMAL)
            self.assertEqual(overview["inspection_marker"].status, STATUS_REVIEW)
            self.assertEqual(overview["substation_led_meter"].status, STATUS_NORMAL)
            self.assertEqual(overview["pump_analog_meters"].status, STATUS_NORMAL)
            self.assertEqual(overview["gas_o2"].status, STATUS_ABNORMAL)
            self.assertEqual(overview["roller_jam"].status, STATUS_ABNORMAL)
            self.assertEqual(overview["equipment_temperature"].status, STATUS_REVIEW)
            self.assertNotIn("station_number", overview)
            thermal_events = [
                item for item in report.details if item.item_id == "roller_jam"
                and item.capture_id
            ]
            self.assertEqual(len(thermal_events), 1)
            self.assertIn("70.00℃", thermal_events[0].result)
            self.assertIn("9 号编号牌", thermal_events[0].result)
            self.assertEqual(thermal_events[0].evidence_path, str(thermal))

            payload = render_report_docx(report)
            document = Document(BytesIO(payload))
            table_text = "\n".join(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("最高温 70.00℃", table_text)
            self.assertIn("9 号牌位置", table_text)
            self.assertEqual(len(document.inline_shapes), 1)

            corrected = result.to_dict()
            corrected["objects"] = []
            repository.save_correction(
                capture_id,
                corrected,
                operator="reviewer",
                reason="现场复核仪表证据不足",
            )
            corrected_report = build_batch_report(repository, "report-batch")
            corrected_overview = {
                item.item_id: item for item in corrected_report.overview
            }
            self.assertEqual(
                corrected_overview["pump_analog_meters"].status,
                STATUS_REVIEW,
            )
            self.assertTrue(
                any("人工修正" in issue for issue in corrected_report.quality_issues)
            )

    def test_docx_contains_required_sections_and_geometry(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            evidence = Path(temp_dir) / "evidence.png"
            evidence.write_bytes(MINIMAL_PNG)
            payload = render_report_docx(build_prototype_report(str(evidence)))
        self.assertTrue(payload.startswith(b"PK"))
        document = Document(BytesIO(payload))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("煤矿实验室 GO2 智能巡检报告", text)
        self.assertIn("1. 项目总览", text)
        self.assertIn("2. 异常与待复核明细", text)
        self.assertIn("6. 复核与签字", text)
        self.assertGreaterEqual(len(document.tables), 7)
        overview = document.tables[1]
        self.assertEqual(len(overview.rows), 14)
        self.assertIn("tblHeader", overview.rows[0]._tr.xml)
        self.assertIn('w:w="9360"', overview._tbl.xml)
        self.assertEqual(len(document.inline_shapes), 1)

        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with ZipFile(BytesIO(payload)) as archive:
            document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
            styles_xml = ElementTree.fromstring(archive.read("word/styles.xml"))
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertEqual(len(media), 1)
            self.assertTrue(archive.read(media[0]))
        page_size = document_xml.find(".//w:sectPr/w:pgSz", namespace)
        margins = document_xml.find(".//w:sectPr/w:pgMar", namespace)
        self.assertEqual(page_size.attrib[f"{{{namespace['w']}}}w"], "12240")
        self.assertEqual(page_size.attrib[f"{{{namespace['w']}}}h"], "15840")
        for side in ("top", "right", "bottom", "left"):
            self.assertEqual(margins.attrib[f"{{{namespace['w']}}}{side}"], "1440")
        for table in document_xml.findall(".//w:tbl", namespace):
            width = table.find("./w:tblPr/w:tblW", namespace)
            indent = table.find("./w:tblPr/w:tblInd", namespace)
            grid = table.findall("./w:tblGrid/w:gridCol", namespace)
            self.assertEqual(width.attrib[f"{{{namespace['w']}}}w"], "9360")
            self.assertEqual(indent.attrib[f"{{{namespace['w']}}}w"], "120")
            self.assertEqual(
                sum(int(column.attrib[f"{{{namespace['w']}}}w"]) for column in grid),
                9360,
            )
        self.assertIsNone(document_xml.find(".//w:trHeight", namespace))
        self.assertIn("PAGE", footer_xml)
        self.assertIn("NUMPAGES", footer_xml)
        self.assertIn("Microsoft YaHei", ElementTree.tostring(styles_xml, encoding="unicode"))


class ReportApiTests(unittest.TestCase):
    def test_report_download_not_ready_and_missing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            inbox = root / "dataset_inbox"
            write_batch(inbox)
            paths = {name: root / f"{name}.json" for name in ("classes", "stations", "modules")}
            paths["classes"].write_text("{}", encoding="utf-8")
            paths["stations"].write_text("{}", encoding="utf-8")
            paths["modules"].write_text(
                json.dumps(
                    {
                        name: {"enabled": False}
                        for name in (
                            "coal_presence",
                            "station_number",
                            "digital_meter",
                            "analog_meter",
                        )
                    }
                ),
                encoding="utf-8",
            )
            settings = root / "app.json"
            settings.write_text(
                json.dumps(
                    {
                        "dataset_inbox_path": str(inbox),
                        "storage_root": str(root / "runtime"),
                        "database_path": str(root / "runtime" / "inspection.db"),
                        "classes_path": str(paths["classes"]),
                        "stations_path": str(paths["stations"]),
                        "modules_path": str(paths["modules"]),
                        "detector": {"backend": "noop", "weights": None},
                    }
                ),
                encoding="utf-8",
            )
            previous = os.environ.get("GO2_INSPECTION_SETTINGS")
            os.environ["GO2_INSPECTION_SETTINGS"] = str(settings)
            try:
                with patch("app.offline_import.OfflineBatchManager.start"):
                    with TestClient(create_app()) as client:
                        batch_id = client.get("/api/v1/offline-batches").json()["items"][0]["batch_id"]
                        queued = client.post(f"/api/v1/offline-batches/{batch_id}/import")
                        self.assertEqual(queued.status_code, 200)
                        response = client.get(
                            f"/api/v1/offline-batches/{batch_id}/report.docx"
                        )
                        self.assertEqual(response.status_code, 409)
                        premature_confirmation = client.post(
                            f"/api/v1/offline-batches/{batch_id}/confirm-report"
                        )
                        self.assertEqual(premature_confirmation.status_code, 409)
                        self.assertEqual(
                            client.get("/api/v1/offline-batches/missing/report.docx").status_code,
                            404,
                        )
            finally:
                if previous is None:
                    os.environ.pop("GO2_INSPECTION_SETTINGS", None)
                else:
                    os.environ["GO2_INSPECTION_SETTINGS"] = previous


if __name__ == "__main__":
    unittest.main()
